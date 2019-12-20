from bs4 import BeautifulSoup
import requests
import xlwt
import os
from docx import Document
import pickle
from progressbar import ProgressBar


def get_content(URL):

    field_title = [
        "Learning Outcome",
        "Assessment",
        "Assessor",
        "Target Year of Achievement",
        "Related Learning Outcomes",
        "Overview",
        "Resources",
    ]
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/78.0.3904.97 Safari/537.36"
    }
    page = requests.get(URL, headers=headers)
    soup = BeautifulSoup(page.content, "html.parser")
    extended_table = soup.find("div", attrs={"class": "content"})
    html_table = soup.find("div", attrs={"class": "pure-u-1 pure-u-md-1 pure-u-lg-18-24"})
    table = soup.find("table", attrs={"class": "lotable"})

    # <h2 inside class="bread" contains the synopsis of the lo>
    synopsis_str = "None"
    try:
        synopsis_str = soup.find("h2").contents[0]
    except OSError:
        synopsis_str = "None Error"
        print(f"There was an error fetching the synopsis of {URL}")
    else:
        pass

    # <th> contains the lo title
    title = []
    for row in table.find_all("th")[1:]:
        title.append(row.get_text())

    # <td> contains field names and content - tecnically this is not needed.
    content = []
    for row in table.find_all("td")[1:]:
        if row.get_text().strip() == None:
            content.append("NONE")
        else:
            content.append(row.get_text())

    # The <strong> contains the field names, the LO is in <strong> as well so we need to remove that from the list.
    headings = []
    for row in table.find_all("strong"):
        headings.append(row.get_text())

    # cleaning content from headings
    for item in field_title:
        content.remove(item)

    # cleaning up content
    clear_content = [item.strip() for item in content]

    # cleaning heading from content
    headings.remove(content[0])

    class Url_content(object):
        def __init__(
            self,
            obj_clear_content=clear_content,
            obj_headings=headings,
            obj_title=title,
            obj_table=html_table,
            obj_ex_table=extended_table,
            obj_synopsis=synopsis_str,

        ):
            self.content = obj_clear_content
            self.headings = obj_headings
            self.title = obj_title
            self.table = obj_table
            self.extended_table = obj_ex_table
            self.synopsis = obj_synopsis
    obj_content = Url_content()
    if len(clear_content) == 8:
        return obj_content
    else:
        return print(
            f"Attention! There was an error while parsing the content of url:{URL}: \nThe table {title} has lengh diferent than 8."
        )


def url_generator(input_file,output_file,exceptions):
    f = open(input_file, "r")
    lo_id = []
    lo_index = []

    for line in f:
        lo_id.append(line.split(" ")[0])
        lo_index.append(int(line.split(" ")[1].strip()))
    f.close()

    # For each lo_id creates append lo_url n = lo_index times with
    lo_url = []
    for id, index in zip(lo_id, lo_index):
        i = 0
        while i != index:
            lo_url.append(f"https://www.rcophth.ac.uk/learningoutcomes/{id}{i + 1}")
            i = i + 1

    exceptions_lo = [item[0] for item in exceptions]
    exceptions_url = [item[1] for item in exceptions]

    print(f"Found the following exceptions: {exceptions_lo}.\nReplacing the url's for {exceptions_url}")

    url_array_index = [item.split("https://www.rcophth.ac.uk/learningoutcomes/")[1] for item in lo_url]
    exception_indices = [key for key, val in enumerate(url_array_index) if val in set(exceptions_lo)]

    for indices, url in zip(exception_indices, exceptions_url):
         lo_url[indices] = url

    new_file = open(output_file, "w+")
    for item in lo_url:
        new_file.write(f"{item}\n")
    new_file.close()
    return lo_url


def create_xls_table(lo_content,path):
    #content, headings, title, number
    book = xlwt.Workbook()
    sheet = book.add_sheet("lo_sheet")

    col = ["A","B","C","D","E","F","G","H","I"]
    fmt = xlwt.Style.easyxf("""font: name Arial;""",)
    heading_array = lo_content[0].headings

    # wrinting the first row - the headings
    sheet.write(0,0,"Title")
    for col in range(1,len(heading_array)):
        sheet.write(0,col,lo_content[0].headings[col-1])

    # now need to write the content of each LO
    row = 1
    for lo in range(len(lo_content)):
        sheet.write(row,0,lo_content[lo].title[0])
        for col in range(1,len(heading_array)):
            sheet.write(row,col,lo_content[lo].content[col-1])
        row = row + 1

    xls_path = path/"lo.xls"
    book.save(xls_path)

    return True

def create_html_table(html_file_name,content,path):
    pwd = path/html_file_name
    new_file = open(pwd, "w+")
    for line in content:
        new_file.write(f"{line}")
    new_file.close()

    return True

def mkdir(pwd,foldername):
    folder_path = pwd / foldername
    try:
        os.mkdir(folder_path)
    except OSError:
        print(f"Creation of the directory {folder_path} failed, probably because it already exists")
    else:
        print(f"Successfully created the directory {folder_path}")

    return True


def create_one_doc(title,synopsis,headings,content):

    document=Document("CA1.docx")
    document.add_heading(f"{content[0]} - {title[0]}", 0)
    document.add_paragraph(synopsis, style='Quote')

    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Atribute'
    hdr_cells[1].text = 'Id'
    for head,attr in zip(headings,content):
        row_cell = table.add_row().cells
        row_cell[0].text = head
        row_cell[1].text = attr

    document.add_page_break()
    document_name = content[0]
    document.save(f"{document_name}.docx")

    return True

def create_doc(lo_title,lo_synopsis,lo_headings,lo_content):

    document = Document("template.docx")
    pbar = ProgressBar()
    document_name = "lo_all"
    for tit,syn,head,cont in pbar(zip(lo_title,lo_synopsis,lo_headings,lo_content)):
        document.add_heading(f"{cont[0]} - {tit[0]}", 0)
        document.add_paragraph(syn, style='Quote')

        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Atribute'
        hdr_cells[1].text = 'Id'
        for h, attr in zip(head, cont):
            row_cell = table.add_row().cells
            row_cell[0].text = h
            row_cell[1].text = attr

        document.add_page_break()

    document.save(f"{document_name}.docx")

    return True

def get_domains(url_domain):

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/78.0.3904.97 Safari/537.36"
    }
    page = requests.get(url_domain, headers=headers)
    soup = BeautifulSoup(page.content, "html.parser")
    domains = [8,1,1,1,1]
    i = 1
    for number in domains:
        for j in range(1,number+1):
            lo_group = soup.find("div", attrs={"class": f"theme-domain-link d{i}m{j}"})
        i = i+1
    return lo_group
